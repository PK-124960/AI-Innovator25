import ollama
import streamlit as st
import json 
import re
import ast 
import csv
import os
import fitz
import qdrant_client
import numpy as np

from sentence_transformers import SentenceTransformer
from numpy.linalg import norm
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from thefuzz import process, fuzz
    THEFUZZ_AVAILABLE = True
    print("✅ 'thefuzz' library is available.")
except ImportError:
    THEFUZZ_AVAILABLE = False
    print("⚠️ 'thefuzz' library not found. Fuzzy matching is disabled.")
    
OLLAMA_HOST = 'http://ollama:11434' 
LLM_MODEL = 'scb10x/llama3.1-typhoon2-8b-instruct:latest'

@st.cache_resource
def init_ollama_client():
    try:
        client = ollama.Client(host=OLLAMA_HOST, timeout=300)
        client.list()
        print("✅ Ollama connection successful.")
        return client, True
    except Exception as e:
        print(f"❌ Ollama connection failed: {e}")
        return None, False

ollama_client, OLLAMA_AVAILABLE = init_ollama_client()

# Page 1 : Draft Generation

PROMPT_TEMPLATES = {
        "กระดาษข่าวร่วม (ทท.)": """
    # ภารกิจ
    คุณคือเสมียนเอกผู้เชี่ยวชาญการร่างหนังสือราชการทหาร ภารกิจของคุณคือแปลงข้อความภาษาพูดของผู้ใช้ให้เป็น "เนื้อความ" ของ "กระดาษข่าวร่วม (ทท.)" ที่มีความถูกต้องตามระเบียบงานสารบรรณ และสอดคล้องกับเจตนารมณ์ของผู้สั่งการ

    # รูปแบบการทำงาน (Pattern to Follow)
    คุณต้องเลียนแบบรูปแบบการทำงานต่อไปนี้อย่างเคร่งครัด: รับข้อความต้นฉบับ แล้วแปลงเป็นผลลัพธ์ที่คาดหวัง

    ---
    ### ตัวอย่างที่ 1: การเชิญประชุม/ขอให้ดำเนินการ ###
    [ข้อความต้นฉบับจากผู้ใช้]: "เราจะจัดประชุมเรื่องการใช้ AI ในวันที่ 15 ส.ค. ที่ห้องประชุมใหญ่ เพื่อให้ทุกคนเข้าใจเทคโนโลยีใหม่ๆ อยากให้แต่ละแผนกส่งคนมาแผนกละ 2 คน ช่วยส่งชื่อภายในวันที่ 10 ส.ค. ด้วยนะ ถามรายละเอียดได้ที่พี่สมชาย เบอร์ 1234"
    [ผลลัพธ์ที่คาดหวัง]:
    ๑. ด้วย [ชื่อหน่วยงานผู้จัด] มีกำหนดจัดประชุมเชิงปฏิบัติการในหัวข้อ "การประยุกต์ใช้ปัญญาประดิษฐ์ในการปฏิบัติงาน" ในวันที่ ๑๕ ส.ค. ๖๗ ณ ห้องประชุม [ชื่อห้องประชุม]
    ๒. การประชุมดังกล่าวมีวัตถุประสงค์เพื่อให้กำลังพลของหน่วยมีความรู้ความเข้าใจเกี่ยวกับเทคโนโลยีปัญญาประดิษฐ์และสามารถนำมาปรับใช้เพื่อเพิ่มประสิทธิภาพในการปฏิบัติงานได้
    ๓. จึงขอให้แต่ละส่วนราชการในสังกัด ได้โปรดพิจารณาจัดส่งกำลังพลเข้าร่วมการประชุมฯ ตามความเหมาะสม จำนวน ๒ นาย และกรุณารวบรวมรายชื่อแจ้งให้ผู้จัดทราบภายในวันที่ ๑๐ ส.ค. ๖๗ เพื่อดำเนินการในส่วนที่เกี่ยวข้องต่อไป
    ๔. หากประสงค์จะสอบถามรายละเอียดเพิ่มเติม สามารถประสานได้ที่ [ยศ ชื่อ-สกุล ผู้ประสานงาน] โทร. [เบอร์โทรศัพท์ผู้ประสานงาน]
    ---
    ### ตัวอย่างที่ 2: การรายงานผล/ขอข้อมูล ###
    [ข้อความต้นฉบับจากผู้ใช้]: "ฝ่ายเราต้องทำรายงานสรุปผลการปฏิบัติงานของปีที่แล้ว อยากจะขอข้อมูลจากทุกแผนกเลย ช่วยส่งข้อมูลผลงานเด่นๆ มาให้ภายในวันที่ 25 ของเดือนนี้นะครับ ส่งมาที่อีเมล Plan.div@rtarf.mi.th"
    [ผลลัพธ์ที่คาดหวัง]:
    ๑. ด้วย [ชื่อหน่วยงานผู้จัด] มีความจำเป็นต้องรวบรวมข้อมูลเพื่อจัดทำรายงานสรุปผลการปฏิบัติงานประจำปีงบประมาณที่ผ่านมา เพื่อใช้เป็นข้อมูลในการวางแผนและพัฒนาองค์กร
    ๒. เพื่อให้การจัดทำรายงานดังกล่าวเป็นไปด้วยความเรียบร้อยและมีข้อมูลที่ครบถ้วนสมบูรณ์ จึงมีความจำเป็นต้องได้รับข้อมูลผลการปฏิบัติงานที่สำคัญจากทุกส่วนราชการในสังกัด
    ๓. จึงขอความร่วมมือมายังส่วนราชการของท่าน ได้โปรดพิจารณาจัดส่งข้อมูลผลการปฏิบัติงานที่สำคัญหรือโครงการที่เป็นที่ประจักษ์ของหน่วย มายัง [ชื่อหน่วยงานผู้จัด] ภายในวันที่ ๒๕ [ระบุเดือนและปี] โดยจัดส่งผ่านทางไปรษณีย์อิเล็กทรอนิกส์ [ระบุอีเมล]
    ๔. หากประสงค์จะสอบถามรายละเอียดเพิ่มเติม สามารถประสานได้ที่ [ยศ ชื่อ-สกุล ผู้ประสานงาน] โทร. [เบอร์โทรศัพท์ผู้ประสานงาน]
    ---

    # คำสั่ง
    จงใช้รูปแบบจากตัวอย่างข้างต้นเพื่อแปลง **[ข้อความต้นฉบับจากผู้ใช้]` ที่จะได้รับต่อไปนี้ ให้เป็น `[ผลลัพธ์ที่คาดหวัง]` ที่สมบูรณ์

    **กฎเหล็ก:**
    1.  ผลลัพธ์สุดท้ายต้องเป็นเนื้อหาของหนังสือราชการที่มี ๔ ข้อเท่านั้น
    2.  ห้ามใส่คำอธิบาย, ห้ามใส่แท็ก, ห้ามใส่หัวข้อ หรือข้อความอื่นใดนอกเหนือจากเนื้อความ ๔ ข้อนั้นโดยเด็ดขาด
    3.  ปรับแก้สำนวนให้เป็นภาษาราชการทหารตาม "ระดับความเป็นทางการที่ต้องการ"
    4.  ข้อมูลที่ไม่มีในต้นฉบับ เช่น ชื่อหน่วยงาน, ชื่อห้องประชุม, ยศ ให้ใช้ตัวยึดตำแหน่ง (Placeholder) เช่น `[ชื่อหน่วยงานผู้จัด]`
    5.  แสดงผลลัพธ์สุดท้ายเท่านั้น ไม่ต้องมีคำว่า [ผลลัพธ์ที่คาดหวัง]

    **ตอนนี้ จงเริ่มทำงานกับข้อมูลจริงที่จะได้รับต่อไปนี้:**
    """,

        # Prompt for "บันทึกข้อความ"
    "บันทึกข้อความ": """
    # ภารกิจ
    คุณคือเสมียนเอกผู้เชี่ยวชาญการร่างหนังสือราชการทหาร ภารกิจของคุณคือแปลงข้อความภาษาพูดให้เป็น "เนื้อความ" ของ "บันทึกข้อความ" ที่สมบูรณ์แบบ

    # ตัวอย่างการทำงาน
    ต่อไปนี้คือตัวอย่างการแปลงข้อความจากภาษาพูดให้เป็นเนื้อหาบันทึกข้อความที่ถูกต้อง คุณต้องศึกษาและเลียนแบบ "ตรรกะ" และ "สไตล์" จากตัวอย่างเหล่านี้

    ---
    ### ตัวอย่างที่ 1: กรณีขึ้นต้นด้วย "เรียน" ###
    [คำขึ้นต้น]: เรียน
    [ข้อความต้นฉบับ]: "สสท.ทร. เขามาขอวิทยากรบรรยายเรื่องข่าวกรองไซเบอร์วันที่ 27 ม.ค. แต่ตอนนี้เรายังไม่มีคนพร้อมเลย เพราะกำลังทำตำราเรื่องนี้กันอยู่ คงช่วยไม่ได้ ให้ช่วยทำหนังสือปฏิเสธไปให้หน่อย"

    #### ผลลัพธ์ที่ถูกต้อง ####
    ๑. ตามที่ สสท.ทร. ขอรับการสนับสนุนวิทยากรบรรยายในหัวข้อ “การประมาณการข่าวกรองในมิติไซเบอร์” ในวันที่ ๒๗ ม.ค. ๖๘
    ๒. ข้อพิจารณา ปัจจุบัน [ชื่อหน่วยงานผู้จัด] กำลังอยู่ในห้วงของการพัฒนาหลักการและจัดทำตำราในเรื่องดังกล่าว จึงยังขาดเจ้าหน้าที่ที่มีความพร้อมในการบรรยายได้อย่างสมบูรณ์ จึงไม่สามารถให้การสนับสนุนวิทยากรตามที่ร้องขอได้ ทั้งนี้ ได้ประสานกับ สสท.ทร. ในเบื้องต้นแล้ว
    ๓. ข้อเสนอ เห็นควรมีหนังสือแจ้งผลการพิจารณาให้ สสท.ทร. ทราบ
    จึงเรียนมาเพื่อโปรดพิจารณา หากเห็นสมควรกรุณาอนุมัติตามข้อ ๓
    ---
    ### ตัวอย่างที่ 2: กรณีขึ้นต้นด้วย "เสนอ" ###
    [คำขึ้นต้น]: เสนอ
    [ข้อความต้นฉบับ]: "ส.อ.หญิง พรรษวลัย อยากลาออกไปเป็นทหารสัญญาบัตรที่ บก.ทท. ตั้งแต่ 16 เม.ย. 68 หน่วยเราดูแล้วก็เห็นด้วยนะ ช่วยเสนอเรื่องให้หน่อย"

    #### ผลลัพธ์ที่ถูกต้อง ####
    ๑. ตามที่ กกล.นซบ.ทหาร ขอความเห็นชอบเรื่องการขอลาออกจากราชการของ ส.อ.หญิง พรรษวลัย เลี่ยมเพ็ชรรัตน์ ตำแหน่ง [ระบุตำแหน่ง] ซึ่งมีความประสงค์ขอลาออกจากราชการเพื่อบรรจุเป็นนายทหารสัญญาบัตร สังกัด กองบัญชาการกองทัพไทย ตั้งแต่ ๑๖ เม.ย. ๖๘ เป็นต้นไป
    ๒. [ชื่อหน่วยงานผู้จัด] ในฐานะหัวหน้าสายวิทยาการ ได้พิจารณาแล้วเห็นว่าการลาออกดังกล่าวไม่ส่งผลกระทบต่อการปฏิบัติงานของหน่วย และเพื่อเป็นขวัญกำลังใจและความก้าวหน้าในสายอาชีพของกำลังพล จึงเห็นชอบการลาออกดังกล่าว
    ๓. ข้อเสนอ เห็นควรแจ้งผลการพิจารณาให้หน่วยต้นสังกัดของกำลังพลทราบ เพื่อดำเนินการในส่วนที่เกี่ยวข้องต่อไป
    จึงเสนอมาเพื่อกรุณาทราบและดำเนินการต่อไป
    ---

    ## 📝 งานของคุณ (Your Task) ##
    ต่อไปนี้คือข้อมูลสำหรับงานจริง คุณต้องสร้างเฉพาะเนื้อความของบันทึกข้อความตามรูปแบบของ `#### ผลลัพธ์ที่ถูกต้อง ####` ในตัวอย่าง โดยยึดถือกฎต่อไปนี้อย่างเคร่งครัด:

    1.  **กฎคำลงท้าย:** ต้องสอดคล้องกับ `[คำขึ้นต้น]` ที่ได้รับมา
        - หากเป็น **เรียน**, ให้ลงท้ายด้วย **"จึงเรียนมาเพื่อโปรด..."**
        - หากเป็น **เสนอ**, ให้ลงท้ายด้วย **"จึงเสนอมาเพื่อกรุณา..."**
    2.  **กฎโครงสร้าง:** ต้องมี ๓ ข้อ ตามตรรกะ "เหตุ -> พิจารณา -> ข้อเสนอ" และข้อ ๓ ต้องขึ้นต้นด้วย "ข้อเสนอ"
    3.  **กฎเนื้อหา:** ห้ามสร้างข้อมูลที่ไม่มีในต้นฉบับโดยเด็ดขาด
    4.  **กฎการแสดงผล (สำคัญที่สุด):** ห้ามตอบกลับสิ่งอื่นใดนอกจากเนื้อความที่สมบูรณ์เท่านั้น **จงเริ่มต้นคำตอบของคุณด้วย "๑." ทันที**
    """
}
    
def draft_generation(client, user_prompt: str, doc_type: str, formality_level: str, doc_salutation: str = ""):

    if not OLLAMA_AVAILABLE:
        return "ระบบ AI ไม่พร้อมใช้งาน"

    system_prompt = PROMPT_TEMPLATES.get(doc_type)

    if not system_prompt:
        return f"เกิดข้อผิดพลาด: ไม่พบ Prompt Template สำหรับประเภทเอกสาร '{doc_type}'"
    
    full_user_content = f"""
    [ข้อความต้นฉบับจากผู้ใช้]
    ---
    {user_prompt}
    ---
    ระดับความเป็นทางการที่ต้องการ: {formality_level}
    """

    try:
        response = client.chat(
            model=LLM_MODEL,
            messages=[
                {'role': 'system', 'content': system_prompt}, 
                {'role': 'user', 'content': full_user_content}  
            ],
            options = {
                'temperature': 0.05,       
                'num_predict': 2048,       
                'top_p': 0.5,              
                'repetition_penalty': 1.15 
            }
        )
        cleaned_response = response['message']['content'].replace("```", "").strip()
        return cleaned_response
    except Exception as e:
        return f"เกิดข้อผิดพลาดในการเรียกใช้ AI: {e}"

    
# Page 2 : Outgoing Letter Creation

def extract_structured_data(client, ocr_text_content: str, document_type: str, system_prompt: str, user_prompt_template: str):
    """Calls LLM to extract structured data from OCR text using the pre-configured client."""
    if not OLLAMA_AVAILABLE or client is None:
        raise ConnectionError("Ollama client is not available for data extraction.")
    
    if not ocr_text_content or ocr_text_content.isspace():
        raise ValueError("No OCR text provided for data extraction.")
    
    user_prompt = user_prompt_template.format(ocr_text=ocr_text_content)
    try:
        response = client.chat(
            model=LLM_MODEL,
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            format="json",
            options={'temperature': 0.05, 'num_predict': 3500}
        )
        content_str = response['message']['content']
        
        print("="*50)
        print("RAW RESPONSE FROM OLLAMA:")
        print(content_str)
        print("="*50)

        try:
            return json.loads(content_str)
        except json.JSONDecodeError:
            json_match = re.search(r'\{.*\}', content_str, re.DOTALL)
            if json_match:
                return json.loads(json_match.group(0))
            else:
                raise ValueError("LLM did not return a valid JSON object.")

    except Exception as e:
        print(f"Error during Ollama chat for data extraction: {e}")
        raise e

        
CORRECT_UNIT_ABBREVIATIONS = [
    "กธก.ศซบ.ทหาร", "กวก.ศซบ.ทหาร", "ผงป.นซบ.ทหาร", "ผกง.นซบ.ทหาร", "นตส.นซบ.ทหาร", "นธน.นซบ.ทหาร", "กกล.นซบ.ทหาร","กขซ.นซบ.ทหาร","กยก.นซบ.ทหาร", 
    "กตซ.นซบ.ทหาร", "สปก.นซบ.ทหาร", "กปก.๑ สปก.นซบ.ทหาร", "กปก.๒ สปก.นซบ.ทหาร", "กปก.๓ สปก.นซบ.ทหาร", "กปก.ศซบ.ทหาร", "ศซล.นซบ.ทหาร",
    "กวก.ศซล.นซบ.ทหาร", "รร.ซบ.ทหาร ศซล.นซบ.ทหาร", "กศษ.รร.ซบ.ทหาร ศซล.นซบ.ทหาร", "สน.บก.บก.ทท.", "สลก.บก.ทท.", "สจร.ทหาร", "สตน.ทหาร", "สสก.ทหาร",
    "สสก.บก.ทท.", "สยย.ทหาร", "ลชท.รอง", "ศปร.", "ศซบ.ทหาร", "สธน.ทหาร", "ขว.ทหาร", "ยก.ทหาร", "กบ.ทหาร", "กร.ทหาร", "กน.ทหาร", "สส.ทหาร", "กปท.ศทส.สส.ทหาร",
    "กทค.ศทท.สส.ทหาร", "พัน.ปสอ.สส.ทหาร", "ร้อย.บก.พัน.ส.บก.ทท.สส.ทหาร", "ร้อย.บก.พัน.ส.", "สปช.ทหาร", "นทพ.", "ศรภ.", "ศตก.", "สบ.ทหาร", "กง.ทหาร", "ผท.ทหาร",
    "ยบ.ทหาร", "สนพ.ยบ.ทหาร", "ชด.ทหาร", "บก.สปท.", "วปอ.สปท.", "วสท.สปท.", "สจว.สปท.", "ศศย.สปท.", "สศท.สปท.", "รร.ตท.สปท.", "รร.ชท.สปท.",
    "รอง ผอ.กภศ.ศศย.สปท.", "รอง ผอ.กกว.วสท.สปท.", "สน.พน.วสท.สปท.", "สน.พน.สปท.", "สน.รอง ผบ.สปท.", "สน.เสธ.สปท.", "กพ.ทหาร", "รอง จก.กพ.ทหาร",
    "จก.กพ.ทหาร", "บก.ทท.", "ผช.ผอ.กรภ.ศซบ.ทหาร", "หก.กธก.ศซบ.ทหาร", "กรภ.ศซบ.ทหาร", "กปก.ศซบ.ทหาร", "กวก.ศซบ.ทหาร", "นขต.ศซบ.ทหาร", "ผอ.ศซบ.ทหาร",
    "กยก.ศซบ.ทหาร", "ศซล.นซบ.ทหาร", "ผบ.นซบ.ทหาร", "ผอ.ศซล.นซบ.ทหาร", "ผอ.กรภ.ศซบ.ทหาร", "ผอ.กวก.ศซบ.ทหาร", "ศชป.ทหาร", "ผอ.กวก.ศซบ.ทหาร", "สสจ.ทหาร",
    "สนย.ทหาร", "วสส.สปท.", "สคท.สปท.", "กมศ.บก.สปท.", "รอง เสธ.สปท.", "ผบ.สปท.", "คทส.บก.ทหาร", "ผอ.กนผ.สผอ.สส.ทหาร", "ผอ.สผอ.สส.ทหาร", "จก.สส.ทหาร",
    "ผอ.กศช.สศท.สปท.", "เสธ.สปท.", "กวก.ศชล.นซบ.ทหาร", "สสท.ทร.", "ศชบ.สสท.ทร.", "จก.สสท.ทร.", "จก.สน.ทหาร", "สนพ.กพ.ทหาร","ศซล.นซบ.ทหาร"
]

# Key: wrong word (frequently)
# Value: correct word
OCR_CORRECTION_MAP = {
    "นศ.สรท.":"นศ.สธท.", "ศชบ":"ศซบ", "กวถ.ศชบ.ทหาร": "กวก.ศซบ.ทหาร", "กรก.ศชบ.ทหาร": "กรก.ศซบ.ทหาร",
    "กธก.ศชบ.ทหาร": "กธก.ศซบ.ทหาร", "กวก.ศชบ.ทหาร": "กวก.ศซบ.ทหาร", "กหค.ศทท.สส.ทหาร": "กทค.ศทท.สส.ทหาร",
    "กปภ.ศชบ.ทหาร": "กปก.ศซบ.ทหาร", "กก.กธก.ศชบ.ทหาร": "หก.กธก.ศซบ.ทหาร", "กวภ.ศชบ.ทหาร": "กวก.ศซบ.ทหาร",
    "ศช.ทหาร. ": "ศซบ.ทหาร ", "คุณท.๖๗": "คกนท.๖๗", "สน.พน.วสท.สปท.": "สน.ผบ.วสท.สปท.",
    "สน.พบ.สปท.": "สน.ผบ.สปท.", "กวต.ศชบ.ทหาร": "กวก.ศซบ.ทหาร", "นชต.ศชบ.ทหาร": "นขต.ศซบ.ทหาร",
    "ผอ.ศชบ.ทหาร": "ผอ.ศซบ.ทหาร", "ศชย.สปท. ": "ศศย.สปท. ", "รอง ผอ.กพศ.ศชย.สปท.": "รอง ผอ.กภศ.ศศย.สปท.",
    "สบ.บก.ทท. ": "สน.บก.บก.ทท. ", "ยน.ทหาร": "ยบ.ทหาร", "บก.ทหาร": "บก.ทท.", "สสค.บก.ทท.": "สลก.บก.ทท.",
    "สสภ.ทหาร": "สสก.ทหาร", "ชว.ทหาร": "ขว.ทหาร", "นทฟ. ": "นทพ.", "กวภ.ศช.น.ทหาร": "กวก.ศซบ.ทหาร",
    "ศช.บ.ทหาร": "ศซบ.ทหาร", "กกล.นชช.ทหาร": "กกล.นซบ.ทหาร", "นชช.ทหาร": "นซบ.ทหาร",
    "ศชล.นชช.ทหาร": "ศซล.นซบ.ทหาร", "กปช.ศชบ.สสท.ทร. ": "กปซ.ศซบ.สสท.ทร.","ถวก.ศชล.นซบ.ทหาร": "กวก.ศซล.นซบ.ทหาร", 
    "กศช.สศท.สปท.": "กศษ.สศท.สปท.", "เสร.สปท.": "เสธ.สปท.", "นชบ.ทหาร": "นซบ.ทหาร",
    "รธ.ชน.ทหาร": "รร.ซบ.ทหาร", "สน.ทหาร": "สบ.ทหาร", "กสม.สน.ทหาร. ": "กสบ.สบ.ทหาร. ",
    "นชต.ศช.ทหาร": "นขต.นซบ.ทหาร", "กวจ.ศชน.ทหาร": "กวก.ศซบ.ทหาร", "ผอ.ศช.ปทหาร": "ผอ.ศซบ.ทหาร",
    "กน.ทหาร": "กบ.ทหาร", "ศตถ. ": "ศตก. ", "สคท.สปท.": "สศท.สปท.", "กรภ.ศชบ.ทหาร": "กธก.ศซบ.ทหาร",
    "รร.รปภ.ศธ. ": "รร.รปภ.ศรภ.", "นทท.": "นทพ.", "กรมทหาร": "กร.ทหาร", "คชช.ทหาร": "ศซบ.ทหาร",
    "ถนนผจงพหาร": "กนผ.กร.ทหาร", "สวผ.ยก.ทหาร": "สวฝ.ยก.ทหาร", "กหศ.ศสภ.ยก.ทหาร": "กฝศ.ศสภ.ยก.ทหาร",
    "กหม.นก.สปท.": "กทด.บก.สปท.", "เลขา.สปท.": "เสธ.สปท.", "ผอ.บทว.สปท.": "ผอ.บฑว.สปท.",
    "กจก.สนส. กม.ทหาร": "กจก.สบส.กบ.ทหาร", "กสม.สน.ทหาร": "กสบ.สบ.ทหาร", "กพศ.ศสภ.ยก.ทหาร": "กฝศ.ศสภ.ยก.ทหาร",
    "ถนนผ.กร.ทหาร": "กนผ.กร.ทหาร", "ศชบ.ทอ.": "ศซบ.ทอ.", "รร.รปภ.ศธ.": "รร.รปภ.ศรภ.", "กคช.บก.นทพ.": "กกช.บก.นทพ.",
    "กบ.สคร.กร.ทหาร": "กบภ.สกร.กร.ทหาร", "กห.อต๊อด.๑๐.๑":"กห ๐๓๐๑.๑๐.๑","จึงเสนอมามาเพื่อกรุณาพิจารณา":"จึงเสนอมาเพื่อกรุณาพิจารณา",
    "๕๗๖๓๙(๔๗).":"๕๗๒๑๗๔๗).","๐-๒๕๗๒.๑๗๔๗.":"๐ ๒๕๗๒ ๑๗๔๗","กห.อต๊อก.๑๐.๑":"กห ๐๓๐๑.๑๐.๑","กปภ.๓":"กปก.๓","กธถ.ศชบ.ทหาร":"กธก.ศซบ.ทหาร",
    "กรก.ศชบ.ทหาร":"กธก.ศซบ.ทหาร","ผช.ผอ.กรก.ศชบ.ทหาร":"ผช.ผอ.กรภ.ศซบ.ทหาร","กท.อต๊อก.๑":"กห ๐๓๐๑.๑๐.๑","ผช.ผอ.กรก.ศชบ.ทหาร":"ผช.ผอ.กรภ.ศซบ.ทหาร",
    "กก.กธก.ศซบ.ทหาร":"หก.กธก.ศซบ.ทหาร","กกล.นชบ.ทหาร":"กกล.นซบ.ทหาร","๐.๒๒๗๕.๕๗๑๖":"๐ ๒๒๗๕ ๕๗๑๖","อิลล์":"ฮิลส์","ไม่กำหนดชื่อ":"ไม่กำหนดชั้นยศ",
    "ผอ.กพศ.ศคย.สปท.":"ผอ.กภศ.ศศย.สปท.","ผอ.ศคย.สปท.":"ผอ.ศศย.สปท.", "..สสท.ทร.(ศซบ.โทร.๕๗๘๙)":"สสท.ทร. (ศซบ. โทร.๕๗๘๙๐)",
    "คานฑ์.๖๗":"คกนท.๖๗","สน.พ.วสท.สปท.":"สน.ผบ.วสท.สปท.","สน.ผ.สปท.":"สน.ผบ.สปท.","ผอ.กพ.วสท.สปท.":"ผอ.กพผ.วสท.สปท.",
    "นายทหารอุบมิติข่าว":"นายทหารอนุมัติข่าว","กระดาษเชิญข่าวร่วม (ทท.)":"กระดาษเขียนข่าวร่วม (ทท.)",
    "จึงเสนอมาระบกวนโปรด":"จึงเสนอมาเพื่อโปรด", "ลาอฉก.": "ลาออก", "0":"๐", "1":"๑", "2":"๒", "3":"๓", "4":"๔", "5":"๕", "6":"๖", "7":"๗", "8":"๘", "9":"๙"
}
     

def post_process_ocr_text(ocr_text: str, fuzzy_enabled: bool = False) -> str:
    
    if not ocr_text or not isinstance(ocr_text, str):
        return ""

    processed_text = ocr_text

    sorted_correction_keys = sorted(OCR_CORRECTION_MAP.keys(), key=len, reverse=True)
    for wrong_word_key in sorted_correction_keys:
        correct_word_value = OCR_CORRECTION_MAP[wrong_word_key]
        processed_text = processed_text.replace(wrong_word_key, correct_word_value)

    if fuzzy_enabled and THEFUZZ_AVAILABLE:
#         print("Fuzzy matching is ENABLED.") 
        try:
            words_in_text = re.findall(r'\b[\w.-]+\b', processed_text, re.UNICODE)
            final_words = []
            MIN_FUZZY_SCORE = 80

            for word in words_in_text:
                is_potential_abbreviation = (
                    '.' in word and
                    any(c.isupper() for c in word) and
                    not word.isdigit() and
                    len(word) >= 5 
                )
                if is_potential_abbreviation:

                    best_match_tuple = process.extractOne(word, CORRECT_UNIT_ABBREVIATIONS, scorer=fuzz.WRatio)
                    if best_match_tuple and best_match_tuple[1] >= MIN_FUZZY_SCORE:
                        print(f"Fuzzy corrected '{word}' -> '{best_match_tuple[0]}' (Score: {best_match_tuple[1]})")
                        final_words.append(best_match_tuple[0])
                    else:
                        final_words.append(word)
                else:
                    final_words.append(word)
            processed_text = " ".join(final_words)
        except Exception as e:
            print(f"Warning: An error occurred during fuzzy matching: {e}")
            pass 
#     else:
#         print("Fuzzy matching is DISABLED.") 

    processed_text = re.sub(r'\s+', ' ', processed_text).strip()
    processed_text = re.sub(r'\s+\.', '.', processed_text)
    processed_text = processed_text.replace('\n', ' ')
    processed_text = processed_text.replace("“", "\"").replace("”", "\"")
    processed_text = processed_text.replace("‘", "'").replace("’", "'")
    processed_text = re.sub(r'\-{3,}', '', processed_text)
    processed_text = re.sub(r'\*{2,}', '', processed_text)
    processed_text = re.sub(r'\s- ', '', processed_text)
    processed_text = processed_text.replace('#', '')
    processed_text = processed_text.replace('|', '')

    return processed_text

# --- FIELD DEFINITIONS ---
FIELDS_MEMORANDUM = {
    "department": "ส่วนราชการ",
    "document_number": "ที่ (เลขหนังสือ)",
    "date": "วันที่",
    "subject": "เรื่อง",
    "recipient": "เรียน/เสนอ",
    "reference": "อ้างถึง",
    "attachments": "สิ่งที่ส่งมาด้วย",
    "body_main": "เนื้อความหลัก",
    "proposer_rank_name": "ยศ ชื่อผู้เสนอ",
    "proposer_position": "ตำแหน่งผู้เสนอ",
    "approver_rank_name": "ยศ ชื่อผู้อนุมัติ",
    "approver_position": "ตำแหน่งผู้อนุมัติ"
}

FIELDS_JOINT_NEWS_PAPER = {
    "urgency": "ลำดับความเร่งด่วน",
    "confidentiality": "ชั้นความลับ",
    "datetime_group": "หมู่-วัน-เวลา",
    "page_info": "หน้าที่",
    "originator_ref": "ที่ของผู้ให้ข่าว",
    "from_department": "จาก (หน่วยงานผู้ส่ง)",
    "to_recipient": "ถึง (ผู้รับปฏิบัติ)",
    "info_recipient": "ผู้รับทราบ",
    "body_main": "เนื้อหาข่าว",
    "qr_email": "QR Code/Email",
    "responsible_unit": "หน่วย (ผู้รับผิดชอบ)",
    "phone": "โทรศัพท์",
    # "reporter_signature": "ลายเซ็นผู้เขียนข่าว", # การสกัดลายเซ็นอาจยาก
    "reporter_rank_name_position": "ยศ ชื่อ ตำแหน่งผู้เขียนข่าว",
    # "approver_signature": "ลายเซ็นนายทหารอนุมัติข่าว", # การสกัดลายเซ็นอาจยาก
    "approver_rank_name_position": "ยศ ชื่อ ตำแหน่งนายทหารอนุมัติข่าว"
}

def get_extraction(document_type: str):

    if document_type == "บันทึกข้อความ":

        fields_to_extract_map = FIELDS_MEMORANDUM
        fields_to_extract_map_with_new = fields_to_extract_map.copy()
        
        fields_to_extract_map_with_new["proposer_title_suffix"] = "ส่วนต่อท้ายตำแหน่งผู้เสนอ (เช่น ปฏิบัติหน้าที่, ทำการแทน)"
        fields_to_extract_map_with_new["approver_command"] = "ข้อความคำสั่งของผู้อนุมัติ (เช่น อนุมัติ, ทราบ, เห็นชอบ)"
        fields_to_extract_map_with_new["coordinator_info"] = "ข้อมูลผู้ประสานงาน (ยศ ชื่อ ตำแหน่ง และเบอร์โทร)"
        fields_to_extract_map_with_new["main_intent"] = "เจตนาหลักของหนังสือ (สรุป)"
        fields_to_extract_map_with_new["requested_action_details"] = "รายละเอียดการดำเนินการที่ร้องขอ (สรุป)"
        
        field_descriptions_for_prompt = "\n".join([f"- {desc} (ใช้ key: {key})" for key, desc in fields_to_extract_map_with_new.items()])

        json_example_fields = {
            "department": "สบ.ทหาร (กสบ.สบ.ทหาร โทร.ทหาร ๕๗๒๑๘๒๑)",
            "document_number": "ที่ กห ๐๓๑๒/๒๒๕๙",
            "date": "๑๙ ก.ย. ๖๗",
            "subject": "ขอรับการสนับสนุนวิทยากร",
            "recipient": "เรียน ผอ.ศซบ.ทหาร",
            "reference": None,
            "attachments": [
                "๑. กำหนดการและขอบเขตการบรรยายหลักสูตรนายทหารประทวนอาวุโส บก.ทท. รุ่นที่ ๑๒ ประจำปีงบประมาณ พ.ศ. ๒๕๖๘",
                "๒. แบบกรอกประวัติผู้บรรยายและความต้องการของผู้บรรยาย"
            ],
            "body_main": "๑. สบ.ทหาร กำหนดเปิดหลักสูตรนายทหารประทวนอาวุโส บก.ทท. รุ่นที่ ๑๒ ประจำปีงบประมาณ พ.ศ. ๒๕๖๘ ตั้งแต่วันที่ ๘ ต.ค.-๒๐ ธ.ค. ๖๗ ณ ห้องเรียน ๘๐๓ ชั้น ๘ อาคาร ๙ บก.ทท. ซึ่งในหลักสูตรฯ ได้กำหนดการบรรยายให้ความรู้ในหัวข้อวิชาเกี่ยวกับความรู้ความมั่นคงปลอดภัยทางไซเบอร์ รายละเอียดตามสิ่งที่ส่งมาด้วย ๑\n๒. การดำเนินการเปิดหลักสูตรฯ ตามข้อ ๑ สบ.ทหาร มีความประสงค์ขอรับการสนับสนุนวิทยากรเพื่อบรรยายให้ความรู้แก่ผู้เข้ารับการฝึกอบรมฯ ในวัน เวลา และสถานที่ ดังกล่าว โดยขอความกรุณาส่งแบบกรอกประวัติผู้บรรยายและความต้องการของผู้บรรยาย รายละเอียดตามสิ่งที่ส่งมาด้วย ๒ ถึง สบ.ทหาร ภายในวันที่ ๒๐ ก.ย. ๖๗",
            "proposer_rank_name": "พล.ท. (จิรศักดิ์ พรรังสฤษฎ์)",
            "proposer_position": "จก.สบ.ทหาร",
            "proposer_title_suffix": None,
            "approver_rank_name": None,
            "approver_position": None,
            "approver_command": None,
            "coordinator_info": "พ.ต. อานนท์ ดามาพงศ์ ประจำแผนกจัดการศึกษาและวิชาการ กสบ.สบ.ทหาร โทร.ทหาร ๕๗๒๑๘๒๑",
            "main_intent": "ขอรับการสนับสนุนวิทยากร",
            "requested_action_details": "ขอวิทยากรบรรยายหัวข้อความมั่นคงปลอดภัยทางไซเบอร์ และส่งประวัติผู้บรรยายภายใน ๒๐ ก.ย. ๖๗"
        }
        
        system_prompt_extraction = f"""คุณคือ AI ผู้เชี่ยวชาญการสกัดข้อมูลจาก "บันทึกข้อความ" ของราชการไทย ภารกิจของคุณคือการอ่านเนื้อหาจาก OCR อย่างละเอียดและสกัดข้อมูลตามคำแนะนำทีละขั้นตอน (Step-by-Step) เพื่อให้ได้ผลลัพธ์ที่แม่นยำที่สุด จากนั้นให้ตอบกลับเป็น JSON object ที่สมบูรณ์เท่านั้น

        **คำแนะนำและขั้นตอนการสกัดข้อมูล (สำคัญมาก!):**

        **ขั้นตอนที่ 1: สกัดข้อมูลส่วนหัว (Header)**
        - `department`, `document_number`, `date`, `subject`, `recipient`: สกัดข้อมูลตามชื่อหัวข้อ
        - `reference`, `attachments`: สกัดข้อมูลส่วน "อ้างถึง" และ "สิ่งที่ส่งมาด้วย" **ให้ผลลัพธ์เป็น list ของ string เสมอ** แม้จะมีเพียงรายการเดียวก็ตาม หากไม่มีข้อมูล ให้เป็น `null` หรือ list ว่าง `[]`

        **ขั้นตอนที่ 2: สกัดเนื้อหาหลักและผู้ประสานงาน (Body & Coordinator)**
        - `body_main`: 'เนื้อความหลัก' ทั้งหมด ตั้งแต่ข้อ ๑. (หรือย่อหน้าแรก) ไปจนถึงข้อสุดท้ายของเนื้อหา **ก่อน** ประโยคลงท้าย (เช่น 'จึงเรียนมาเพื่อโปรดพิจารณา') หรือ **ก่อน** ข้อมูลผู้ประสานงาน
        - `coordinator_info`: ค้นหาข้อความที่ระบุ "มอบหมายให้..." หรือ "ประสานรายละเอียด..." แล้วสกัดข้อมูลผู้ประสานงานทั้งหมด (ยศ ชื่อ ตำแหน่ง และเบอร์โทร) ออกมาเป็นสตริงเดียว หากไม่พบให้เป็น `null`

        **ขั้นตอนที่ 3: สกัดส่วนผู้ลงนาม (Signatories) อย่างละเอียด**
        - **ผู้เสนอ (Proposer):**
            - `proposer_rank_name`: ยศและชื่อในวงเล็บของผู้เสนอเรื่อง
            - `proposer_position`: ตำแหน่งหลักของผู้เสนอเรื่อง
            - `proposer_title_suffix`: หากมีบรรทัด "ทำการแทน", "รักษาราชการแทน" หรือ "ปฏิบัติหน้าที่" ให้สกัดข้อความนั้นมาใส่ที่นี่ หากไม่มีให้เป็น `null`
        - **ผู้อนุมัติ (Approver):**
            - ค้นหาส่วนที่มีลายเซ็นหรือข้อความเหนือส่วนของผู้เสนอ เช่น "อนุมัติ", "ทราบ", "เห็นชอบ"
            - `approver_command`: สกัดข้อความคำสั่ง เช่น "อนุมัติตามข้อ ๔", "ทราบ"
            - `approver_rank_name`: ยศและชื่อในวงเล็บของผู้อนุมัติ
            - `approver_position`: ตำแหน่งของผู้อนุมัติ

        **ขั้นตอนที่ 4: สกัดข้อมูลเชิงสรุป (Analysis - ทำเป็นขั้นตอนสุดท้าย)**
        - `main_intent`: หลังจากอ่าน `body_main` ทั้งหมดแล้ว ให้สรุป "จุดประสงค์สำคัญที่สุด" ของบันทึกข้อความนี้เป็นวลีสั้นๆ เช่น "ขออนุมัติแก้ไขคำสั่ง", "ขอความเห็นชอบการลาออก", "ขอรับการสนับสนุนวิทยากร"
        - `requested_action_details`: ให้สรุปว่าผู้รับเอกสารต้องทำอะไร (Action Item) และมีกำหนดเวลาเมื่อใด (Deadline) เช่น "พิจารณาอนุมัติการแก้ไขรายชื่อผู้เข้ารับการศึกษา", "ส่งประวัติผู้บรรยายภายใน 20 ก.ย. 67"

        **ฟิลด์ที่ต้องการสกัด (โปรดใช้ key ภาษาอังกฤษ):**
        {field_descriptions_for_prompt}

        ถ้าไม่พบข้อมูลสำหรับฟิลด์ใด ให้ใช้ค่าเป็น `null`
        ตัวอย่าง JSON output ที่คาดหวัง:
        {json.dumps(json_example_fields, indent=2, ensure_ascii=False)} 
                """
    elif document_type == "กระดาษข่าวร่วม (ทท.)":

        fields_to_extract_map = FIELDS_JOINT_NEWS_PAPER
        fields_to_extract_map_with_new = fields_to_extract_map.copy()
        fields_to_extract_map_with_new["coordinator_info"] = "ข้อมูลผู้ประสานงาน (ยศ ชื่อ ตำแหน่ง และเบอร์โทร)"
        fields_to_extract_map_with_new["main_intent"] = "เจตนาหลักของข่าว"
        fields_to_extract_map_with_new["requested_action_details"] = "รายละเอียดการดำเนินการที่ร้องขอจากผู้รับปฏิบัติ"

        field_descriptions_for_prompt = "\n".join([f"- {desc} (ใช้ key: {key})" for key, desc in fields_to_extract_map_with_new.items()])

        json_example_fields = {
            "urgency": "ด่วนมาก",
            "confidentiality": None,
            "datetime_group": "๒๖๐๘๐๐ ก.ค. ๖๗",
            "page_info": "หน้าที่ ๑ ของ ๑ หน้า",
            "originator_ref": "ที่ กห ๐๓๐๒/๔๒๗๓",
            "from_department": "กพ.ทหาร",
            "to_recipient": "สน.บก.บก.ทท. สลก.บก.ทท. สจร.ทหาร สตน.ทหาร สธน.ทหาร สสก.ทหาร สยย.ทหาร ศปร. ศซบ.ทหาร ขว.ทหาร ยก.ทหาร กบ.ทหาร กร.ทหาร สส.ทหาร สปช.ทหาร นทพ. ศรภ. ศตก. สบ.ทหาร กง.ทหาร ผท.ทหาร ยบ.ทหาร ชด.ทหาร และ สปท.",
            "info_recipient": "สนพ.ยบ.ทหาร วปอ.สปท. วสท.สปท. ศศย.สปท. สจว.สปท. รร.ตท.สปท. รร.ชท.สปท. และ สศท.สปท.",
            "body_main": "๑. กพ.ทหาร กำหนดจัดการประชุมเชิงปฏิบัติการ การวิเคราะห์อัตราของส่วนราชการใน บก.ทท. ... ๒. เพื่อให้การดำเนินการตามข้อ ๑ เป็นไปด้วยความเรียบร้อย ขอให้ส่วนราชการจัดผู้แทนเข้าร่วมการประชุมเชิงปฏิบัติการฯ ดังนี้ ...",
            "qr_email": "Kanyarat.y@rtarf.mi.th หรือโทรสาร ๐ ๒๕๗๕ ๖๐๑๕",
            "responsible_unit": "กนผ.สนผพ.กพ.ทหาร",
            "phone": "๐ ๒๕๗๕ ๖๐๑๕",
            "reporter_rank_name_position": "น.อ. (ฐิติพันธ์ บุตรดีสุวรรณ) ผอ.กนผ.สนผพ.กพ.ทหาร",
            "approver_rank_name_position": "พล.ต. (สัมพันธ์ รงศ์จำเริญ) รอง จก.กพ.ทหาร ทำการแทน จก.กพ.ทหาร",
            "coordinator_info": "ร.ท. ภูภวะ สำเร็จผล ร.น. ประจำแผนกนโยบายและแผน กนผ.สนผพ.กพ.ทหาร โทรศัพท์เคลื่อนที่ ๐๘ ๙๑๖๔ ๒๐๒๕ โทร.ทหาร ๕๗๒๑๑๑๘",
            "main_intent": "เชิญเข้าร่วมประชุมเชิงปฏิบัติการฯ และขอให้ส่งรายชื่อผู้แทน",
            "requested_action_details": "จัดผู้แทนเข้าร่วมประชุม (หน.สายวิทยาการและนายทหารกำลังพล หรือสายวิทยาการที่มีแผนปรับเปลี่ยน) และส่งรายชื่อภายในวันศุกร์ที่ ๑๙ ก.ค. ๖๗"
        }

        system_prompt_extraction = f"""คุณคือ AI ผู้เชี่ยวชาญการสกัดข้อมูลจากเอกสาร "กระดาษข่าวร่วม (ทท.)" ของราชการไทย ภารกิจของคุณคือการอ่านเนื้อหาจาก OCR อย่างละเอียดและสกัดข้อมูลตามคำแนะนำทีละขั้นตอน (Step-by-Step) เพื่อให้ได้ผลลัพธ์ที่แม่นยำที่สุด จากนั้นให้ตอบกลับเป็น JSON object ที่สมบูรณ์เท่านั้น

        **คำแนะนำและขั้นตอนการสกัดข้อมูล (สำคัญมาก!):**

        **ขั้นตอนที่ 1: สกัดข้อมูลส่วนหัว (Header)**
        - `urgency`: ความเร่งด่วนของข่าวที่ปรากฏในช่อง 'ลำดับความเร่งด่วน' เช่น 'ด่วนมาก', 'ด่วน', หรือ `null` ถ้าว่าง
        - `datetime_group`: 'หมู่ - วัน - เวลา' ที่อยู่ด้านบนของเอกสาร
        - `originator_ref`: 'ที่ของผู้ให้ข่าว' หรือเลขที่หนังสือของหน่วยงานต้นเรื่อง
        - `from_department`: 'จาก' คือชื่อหน่วยงานที่ส่งข่าวนี้

        **ขั้นตอนที่ 2: สกัดผู้รับ (Recipients)**
        - `to_recipient`: 'ถึง (ผู้รับปฏิบัติ)' ให้รวบรวมรายชื่อหน่วยงานทั้งหมดที่อยู่ภายใต้หัวข้อนี้ แม้จะอยู่คนละบรรทัดก็ตาม ให้รวมเป็นสตริงเดียวคั่นด้วยเว้นวรรค
        - `info_recipient`: 'ผู้รับทราบ' ให้รวบรวมรายชื่อหน่วยงานทั้งหมดภายใต้หัวข้อนี้ รวมเป็นสตริงเดียวเช่นกัน

        **ขั้นตอนที่ 3: สกัดเนื้อหาและผู้ประสานงาน (Body & Coordinator)**
        - `body_main`: 'เนื้อหาข่าว' **ทั้งหมด** ตั้งแต่ข้อ ๑. จนถึงข้อสุดท้ายของเนื้อหา **ก่อน** ส่วนที่จะระบุข้อมูลผู้ประสานงาน หรือก่อนตารางลงนาม
        - `coordinator_info`: ค้นหาข้อความที่ระบุ "รายละเอียดเพิ่มเติมประสาน..." หรือข้อมูลที่คล้ายกัน แล้วสกัดข้อมูลผู้ประสานงานทั้งหมด (ยศ ชื่อ ตำแหน่ง และเบอร์โทร) ออกมาเป็นสตริงเดียว หากไม่พบให้เป็น `null`

        **ขั้นตอนที่ 4: สกัดข้อมูลส่วนท้าย (Footer & Signatories)**
        - `qr_email` / `phone` / `responsible_unit`: สกัดข้อมูลในตารางส่วนท้ายตามชื่อช่อง
        - `reporter_rank_name_position`: ในช่อง "ผู้เขียนข่าว" ให้สกัด ยศ, ชื่อในวงเล็บ, และตำแหน่งทั้งหมด รวมเป็นสตริงเดียว
        - `approver_rank_name_position`: ในช่อง "นายทหารอนุมัติข่าว" ให้สกัด ยศ, ชื่อในวงเล็บ, ตำแหน่ง, และบรรทัด "ทำการแทน/ปฏิบัติหน้าที่แทน" (ถ้ามี) ทั้งหมด รวมเป็นสตริงเดียว

        **ขั้นตอนที่ 5: สกัดข้อมูลเชิงสรุป (Analysis - ทำเป็นขั้นตอนสุดท้าย)**
        - `main_intent`: "เจตนาหลัก" หลังจากอ่าน `body_main` ทั้งหมดแล้ว ให้สรุป "จุดประสงค์สำคัญที่สุด" ของข่าวนี้เป็นวลีสั้นๆ เช่น "เชิญเข้าร่วมประชุม", "ขอให้ส่งข้อมูล", "แจ้งผลการดำเนินการ"
        - `requested_action_details`: "รายละเอียดการดำเนินการที่ร้องขอ" ให้สรุปว่าผู้รับปฏิบัติต้องทำอะไร (Action Item) และมีกำหนดส่งเมื่อใด (Deadline) เช่น "ส่งรายชื่อผู้แทนเข้าร่วมประชุมภายใน ๑๙ ก.ค. ๖๗"

        **ฟิลด์ที่ต้องการสกัด (โปรดใช้ key ภาษาอังกฤษ):**
        {field_descriptions_for_prompt}

        ถ้าไม่พบข้อมูลสำหรับฟิลด์ใด ให้ใช้ค่าเป็น `null` อย่างเคร่งครัด
        ตัวอย่าง JSON output ที่คาดหวัง:
        {json.dumps(json_example_fields, indent=2, ensure_ascii=False)}
                """

    else:
        raise ValueError(f"Document type '{document_type}' is not supported for extraction.")

    user_prompt_template = f"""
    กรุณาสกัดข้อมูลจากเนื้อหาเอกสาร '{document_type}' ต่อไปนี้:
    --- OCR TEXT START ---
    {{ocr_text}}
    --- OCR TEXT END ---
    โปรดตอบกลับเป็น JSON object ที่มีโครงสร้างตามที่ระบุใน system prompt เท่านั้น
    """
    return system_prompt_extraction, user_prompt_template, list(fields_to_extract_map_with_new.keys())
    

def create_docx_from_text(text_content: str, font_name='TH SarabunPSK', font_size=16) -> bytes:

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)

    for line in text_content.split('\n'):
        p = document.add_paragraph(line.strip())
        
        run = p.runs[0] if p.runs else p.add_run('')
        run.font.name = font_name
        run.font.size = Pt(font_size)

        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    from io import BytesIO
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.getvalue()


def log_feedback_to_csv(log_entry: dict, file_path="feedback/feedback_log.csv"):
    try:
        log_dir = os.path.dirname(file_path)
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)
            print(f"Created directory: {log_dir}")

        fieldnames = [
            "timestamp", "document_type", "document_subject",
            "original_text", "edited_text"
        ]
        file_exists = os.path.isfile(file_path)
        
        with open(file_path, 'a', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

            if not file_exists:
                writer.writeheader()
            
            writer.writerow(log_entry)
        
        print(f"Successfully logged feedback to {file_path}")

    except Exception as e:
        print(f"Error logging feedback to CSV: {e}")


def replySec1_generation(client, extracted_info: dict, ocr_text_content: str, num_options: int = 3) -> list:
    if not client:
        print("ERROR [replySec1_generation]: Ollama client is not available.")
        return []
    if not ocr_text_content and not extracted_info:
        print("ERROR [replySec1_generation]: No content or extracted info provided.")
        return []

    sender_department = extracted_info.get("department") or extracted_info.get("from_department") or "[ส่วนราชการต้นเรื่อง]"
    doc_number = extracted_info.get("document_number") or "[เลขที่หนังสือต้นเรื่อง]"
    doc_date = extracted_info.get("date") or "[วันที่หนังสือต้นเรื่อง]"
    doc_subject = extracted_info.get("subject") or "[เรื่องของหนังสือรับ]"
    main_intent = extracted_info.get("main_intent") or "ไม่สามารถระบุเจตนาหลักได้"
    requested_details = extracted_info.get("requested_action_details") or "ไม่พบรายละเอียดการร้องขอ"

    system_prompt = f"""
    <role_definition>
    คุณคือ "เสมียนเอกอัจฉริยะ" แห่งกองบัญชาการกองทัพไทย ผู้มีทักษะขั้นสูงสุดในการวิเคราะห์หนังสือราชการและยกร่าง "ข้อ ๑" ของบันทึกข้อความตอบกลับได้อย่างสมบูรณ์แบบและสละสลวย ภารกิจของคุณคือการสร้างสรรค์ย่อหน้าแรกที่ถูกต้องตามระเบียบงานสารบรรณ โดยอ้างอิงจากข้อมูลของหนังสือต้นเรื่องที่ได้รับมา
    </role_definition>

    <core_instructions>
    วิเคราะห์ข้อมูลทั้งหมดที่ได้รับใน `<input_data>` จากผู้ใช้ จากนั้นให้ทำตามขั้นตอนต่อไปนี้:
    1.  **Analyze Context:** ทำความเข้าใจเจตนาหลักของหนังสือต้นเรื่อง (เช่น ขออนุมัติ, ขอความร่วมมือ, ขอข้อมูล)
    2.  **Match Pattern:** เปรียบเทียบสถานการณ์ที่วิเคราะห์ได้กับ `<case_studies>` ที่ให้ไว้ เพื่อหารูปแบบการร่างที่เหมาะสมที่สุด
    3.  **Generate Options:** สร้าง "ข้อ ๑" ของหนังสือตอบกลับมาทั้งหมด {num_options} รูปแบบ โดยแต่ละรูปแบบอาจปรับเปลี่ยนสำนวนหรือระดับความละเอียดเล็กน้อย แต่ยังคงแก่นของเรื่องไว้
    4.  **Format Output:** จัดรูปแบบผลลัพธ์ให้อยู่ในรูปแบบ JSON Object ที่ถูกต้องตามที่ระบุใน `<output_format_rules>` อย่างเคร่งครัด
    </core_instructions>

    <case_studies>
    จงศึกษาและทำความเข้าใจตรรกะเบื้องหลังกรณีศึกษาเหล่านี้ เพื่อใช้เป็นต้นแบบในการร่าง

    <case name="การอนุมัติและส่งต่อ (Approval & Forwarding)">
      <situation>หน่วยงานย่อยในสังกัด (เช่น กกล.นซบ.ทหาร) ส่งเรื่องมาเพื่อขออนุมัติหรือขอความเห็นชอบในเรื่องที่ต้องให้หน่วยเหนือพิจารณา (เช่น การลาออก, การแต่งตั้ง, การปรับเปลี่ยนตำแหน่ง)</situation>
      <rationale>ย่อหน้าแรกต้องอ้างถึงหนังสือของหน่วยงานย่อยนั้น และสรุปใจความสำคัญของเรื่องที่ขออนุมัติ/เห็นชอบให้ชัดเจน โดยระบุรายละเอียดสำคัญ เช่น ชื่อบุคคล, ตำแหน่ง, และสาเหตุ</rationale>
      <master_example>๑. ตามที่ กธก.ศซบ.ทหาร ขอความเห็นชอบการลาออกจากราชการของ จ.ส.อ. บุรเวทย์ เรื่อศรีจันทร์ ซึ่งมีความประสงค์ขอลาออกจากราชการเพื่อไปปฏิบัติงานบริษัทเอกชน ตั้งแต่ ๑ ก.ค. ๖๗ เป็นต้นไป รายละเอียดตามหนังสือที่อ้างถึงนั้น</master_example>
    </case>

    <case name="การให้ความเห็นชอบภายใน (Internal Endorsement)">
      <situation>หน่วยงานภายนอก (เช่น สบ.ทหาร, วสท.สปท.) ส่งหนังสือมาเพื่อขอรับการสนับสนุน, ขอข้อมูล, หรือขอความร่วมมือในเรื่องต่างๆ (เช่น ขอวิทยากร, ขอข้อมูลวิจัย, ขอให้เข้าร่วมประชุม)</situation>
      <rationale>ย่อหน้าแรกต้องเริ่มต้นด้วยการอ้างถึงหนังสือของหน่วยงานนั้นๆ และสรุป "คำร้องขอ" หลักของพวกเขาให้กระชับและได้ใจความ เพื่อเป็นการทวนเรื่องก่อนจะตอบกลับในข้อถัดไป</rationale>
      <master_example>๑. ตามที่ สบ.ทหาร ขอรับการสนับสนุนวิทยากรบรรยายในหัวข้อวิชา “ความรู้ความมั่นคงปลอดภัยทางไซเบอร์” ในการฝึกอบรมหลักสูตรนายทหารประทวนอาวุโส บก.ทท. รุ่นที่ ๑๒ รายละเอียดตามสิ่งที่ส่งมาด้วยนั้น</master_example>
    </case>

    <case name="การแจ้งเรื่องหรือประกาศให้ทราบโดยทั่วไป">
      <situation>เป็นการแจ้งข่าว, ประกาศ, หรือคำสั่งจากหน่วยเหนือให้หน่วยรองรับทราบเพื่อปฏิบัติ โดยไม่มีการร้องขอจากหน่วยรองมาก่อน (เช่น การจัดประชุม, การสำรวจข้อมูล, การแจ้งนโยบาย)</situation>
      <rationale>ย่อหน้าแรกมักจะขึ้นต้นด้วย "ด้วย..." เพื่อแจ้งถึงเหตุการณ์หรือความจำเป็นที่ทำให้ต้องมีหนังสือนี้ออกมา โดยระบุถึงกิจกรรมหลัก วันที่ และสถานที่ (ถ้ามี)</rationale>
      <master_example>๑. ด้วย กพ.ทหาร กำหนดจัดการประชุมเชิงปฏิบัติการ การวิเคราะห์อัตราของส่วนราชการใน บก.ทท. ในวันที่ ๒๔ ก.ค. ๖๗ ณ ห้องประชุม กพ.ทหาร เพื่อให้การปรับปรุงโครงสร้างเป็นไปด้วยความเรียบร้อย</master_example>
    </case>

    </case_studies>

    "<output_format_rules>\n"
    "**กฎเหล็ก! ต้องปฏิบัติตามอย่างเคร่งครัด:**\n"
    "1.  **JSON Object Only:** ผลลัพธ์สุดท้ายต้องเป็น JSON Object ที่สมบูรณ์แบบเท่านั้น...\n"
    "2.  **Strict Keys:** JSON Object ต้องมี Key เป็น `style_1`, `style_2`, และ `style_3` เท่านั้น\n"
    "3.  **Valid Content:** ค่า (value) ของแต่ละ Key ต้องเป็น String ที่ขึ้นต้นด้วย \"๑. \" และเป็นภาษาราชการที่ถูกต้อง\n"
    "4.  **Analyze and Generate:** วิเคราะห์สถานการณ์, จับคู่กับกรณีศึกษา, แล้วสร้าง {num_options} ตัวเลือกตามรูปแบบ\n" 
    "5.  **No Extraneous Text:** ห้ามมีข้อความใดๆ ปรากฏนอกวงเล็บปีกกาของ JSON object โดยเด็ดขาด\n"
    "</output_format_rules>"

    <example_of_correct_output>
    ```json
    {{
        "style_1": "๑. ด้วย กธก.ศซบ.ทหาร มีความประสงค์ขอแต่งตั้ง ลชท.รอง ให้กับข้าราชการ จำนวน ๖ นาย เนื่องจากได้สำเร็จการฝึกอบรมหลักสูตรตามแนวทางรับราชการของสายวิทยาการความมั่นคงปลอดภัยทางไซเบอร์ เป็นที่เรียบร้อย รายละเอียดตามสิ่งที่ส่งมาด้วย ๑ และ ๒",
        "style_2": "๑. ตามอ้างถึง ผบ.ทสส. ได้กรุณาอนุมัติแต่งตั้งคณะกรรมการพิจารณาผลการบริหารจัดการกำลังพลด้วยสายวิทยาการและเลขหมายความชำนาญการทหารของ บก.ทท. เพื่อดำเนินการแก้ไขปัญหา ในภาพรวมที่มีความซับซ้อนและให้การบริหารจัดการกำลังพลด้วยสายวิทยาการ และเลขหมายความชำนาญการทหารของ บก.ทท. เป็นไปอย่างมีประสิทธิภาพ",
        "style_3": "๑. ตามที่ วสท.สปท. ได้ขอรับความคิดเห็นและข้อเสนอแนะเกี่ยวกับสาระการวิจัยสำหรับนักศึกษาฯ รุ่นต่อไป รายละเอียดปรากฏตามหนังสือที่อ้างถึงนั้น"
    }}
    </example_of_correct_output>
    
    <rules>
        1.  **JSON Object เท่านั้น:** ผลลัพธ์สุดท้ายต้องเป็น JSON Object ที่มี Key เป็น style_1, style_2, style_3 เท่านั้น
        2.  **ห้ามมีข้อความอื่น:** ห้ามมีคำอธิบายหรือข้อความใดๆ นอกเหนือจาก JSON Object โดยเด็ดขาด
        3.  **วิเคราะห์สถานการณ์:** อ่านเนื้อหาทั้งหมดเพื่อจับคู่กับกรณีศึกษาที่ใกล้เคียงที่สุด
        4.  **สร้าง 3 ตัวเลือก:** ยกร่าง "ข้อ ๑" มา {num_options} รูปแบบ โดยอาจดัดแปลงถ้อยคำหรือระดับของรายละเอียดเล็กน้อย แต่ยังคงยึดตามรูปแบบของกรณีศึกษาหลักที่เลือกไว้
        5.  **รูปแบบผลลัพธ์:** แต่ละสตริงต้องขึ้นต้นด้วย "๑. " และใช้ภาษาราชการที่ถูกต้องและสละสลวย
        </rules>
        """    

    user_prompt = f"""
    <input_data>
        <reference_document>
            <sender_department>{sender_department}</sender_department>
            <document_number>{doc_number}</document_number>
            <document_date>{doc_date}</document_date>
            <subject>{doc_subject}</subject>
            <main_intent>{main_intent}</main_intent>
            <requested_details>{requested_details}</requested_details>
        </reference_document>
        <full_ocr_content>
            {ocr_text_content}
        </full_ocr_content>
    </input_data>

    <instruction>
        โปรดวิเคราะห์ข้อมูลทั้งหมดใน <input_data> จับคู่กับกรณีศึกษาที่เหมาะสมที่สุด แล้วยกร่าง "ข้อ ๑" ของบันทึกข้อความตอบกลับมา **{num_options} รูปแบบ** ตามกฎและรูปแบบที่กำหนดใน system prompt อย่างเคร่งครัด
    </instruction>
    """

    try:
        response = client.chat(
            model=LLM_MODEL,
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            format="json",
            options={
                'temperature': 0.6,
                'num_predict': 2048,
                'top_p': 0.9,
                'repetition_penalty': 1.1
            }
        )
        content_str = response['message']['content'].strip()
        print(f"DEBUG [replySec1_generation] Raw LLM response:\n---\n{content_str}\n---")

        try:
            data = json.loads(content_str)
            if isinstance(data, list) and all(isinstance(item, str) for item in data):
                options = [opt.strip() for opt in data if opt.strip()][:num_options]
                if options:
                    print("SUCCESS: Parsed as JSON list.")
                    return options
            elif isinstance(data, dict):
                options = [str(v).strip() for v in data.values() if str(v).strip()][:num_options]
                if options:
                    print("SUCCESS: Parsed as JSON dictionary, extracted values.")
                    return options
        except json.JSONDecodeError:
            print("INFO: Response is not a valid JSON. Trying fallback methods.")

        if content_str.count("๑.") > 1:
            parts = [part.strip() for part in content_str.split("๑.") if part.strip()]
            options = [f"๑. {part}" for part in parts][:num_options]
            if options:
                print("SUCCESS: Fallback - Split by '๑.'")
                return options

        if content_str.startswith('[') and content_str.endswith(']'):
            try:
                evaluated_data = ast.literal_eval(content_str)
                if isinstance(evaluated_data, list):
                    options = [str(item).strip() for item in evaluated_data if str(item).strip()][:num_options]
                    if options:
                        print("SUCCESS: Fallback - Parsed string as a literal list.")
                        return options
            except (ValueError, SyntaxError):
                pass 

        print("WARN: All parsing methods failed. Returning empty list.")
        return []

    except Exception as e:
        print(f"ERROR [replySec1_generation]: Exception during LLM call or processing: {e}")
        return []
        


def _build_prompts_for_intent(reply_intent: str, extracted_info: dict, relevant_internal_data: dict) -> tuple[str, str]:
    
    our_department_name = (relevant_internal_data or {}).get("our_department_name", "[ชื่อหน่วยงานของท่าน]")
    user_provided_opening = extracted_info.pop("user_provided_opening_paragraph", "")

    # --- System Prompt ---
    system_prompt_base = f"""คุณคือ "เสมียนเอกอัจฉริยะ" ผู้เชี่ยวชาญการร่างหนังสือราชการตอบกลับของไทย ภารกิจของคุณคือวิเคราะห์ข้อมูลทั้งหมดที่ได้รับ แล้วยกร่างเนื้อความ (ข้อ ๒, ๓, ...) ต่อจาก "ข้อ ๑" ที่ผู้ใช้กำหนดมาให้สมบูรณ์

<หลักการร่างและการใช้ข้อมูล>
1.  **วิเคราะห์ข้อมูลทั้งหมด:** จงทำความเข้าใจข้อมูลจาก "หนังสือต้นเรื่อง" และ "ข้อมูลสำหรับการตอบกลับ" อย่างละเอียด
2.  **ใช้ข้อมูลให้ครบถ้วน:** ต้องนำข้อมูลจาก "หนังสือต้นเรื่อง" (เช่น เรื่อง, วันที่, สถานที่, รายละเอียดกิจกรรม, ชื่อบุคคล) มาใช้ในการร่างเนื้อหาส่วนที่เหลือ เพื่อให้เนื้อหามีความเชื่อมโยงและสมเหตุสมผล
3.  **สร้างสรรค์อย่างมืออาชีพ:** สามารถเรียบเรียงถ้อยคำและย่อหน้าเพิ่มเติมได้ตามความเหมาะสม เพื่อให้เอกสารมีความสมบูรณ์ สละสลวย และอ่านเข้าใจง่ายเหมือนที่มนุษย์เขียน แต่ต้องอยู่บนพื้นฐานของข้อมูลที่ได้รับเท่านั้น ห้ามสร้างข้อมูลขึ้นมาเอง
4.  **ยึดตามเจตนา:** เนื้อหาที่ร่างต้องสอดคล้องกับ "เจตนาการตอบกลับ" ที่ระบุไว้อย่างชัดเจนตาม <หลักการร่างสำหรับเจตนา>

<กฎการแสดงผล (สำคัญที่สุด)>
- **จงร่างเนื้อหาต่อจาก 'ข้อ ๑' ที่ได้รับมา โดยเริ่มต้นคำตอบของคุณที่ '๒.' ทันที**
- **ห้ามใส่ 'ข้อ ๑' ซ้ำเข้ามาในผลลัพธ์โดยเด็ดขาด**
- ผลลัพธ์ต้องเป็นข้อความธรรมดา (Plain Text) เท่านั้น
- ต้องขึ้นต้นแต่ละย่อหน้าด้วยหมายเลขข้อแบบไทยและจุด (เช่น ๒., ๓., ๔.) โดยไม่มีคำว่า "ข้อ" นำหน้า
- ผลลัพธ์ต้องมีมากสุดได้แค่ 4 ข้อเท่านั้น ห้ามเกิน
- ต้องมีคำลงท้ายที่เหมาะสมกับเจตนา (เช่น จึงเรียนมาเพื่อโปรดพิจารณา, จึงเรียนมาเพื่อโปรดทราบ)
"""

    user_prompt_parts = []
    system_prompt_addendum = "" 

    if reply_intent == "อนุมัติ/เห็นชอบ":
        system_prompt_addendum = """
<หลักการร่างสำหรับเจตนา "อนุมัติ/เห็นชอบ">
- **๒.** ให้อธิบายว่าหน่วยงานของเรา ({our_department_name}) ได้พิจารณาเรื่องที่เสนอมาแล้ว และเห็นว่าสอดคล้องกับภารกิจ หรือเป็นประโยชน์ หรือไม่มีข้อขัดข้อง
- **๓.** ให้ระบุ "ข้อเสนอ" โดยเสนอเพื่อ "อนุมัติ/เห็นชอบ" ในสิ่งที่ร้องขอ และอาจตามด้วยการเสนอให้ "มีหนังสือแจ้ง..." หรือ "ประสานงาน..." ต่อไป
</หลักการร่างสำหรับเจตนา>
"""
    elif reply_intent == "ปฏิเสธ/ไม่เห็นชอบ":
        system_prompt_addendum = """
<หลักการร่างสำหรับเจตนา "ปฏิเสธ/ไม่เห็นชอบ">
- **๒.** ให้อธิบายเหตุผลในการปฏิเสธอย่างสุภาพ (เช่น ติดภารกิจเร่งด่วน, บุคลากรไม่เพียงพอ) และอาจเสริมว่าได้ประสานงานแจ้งเบื้องต้นแล้ว
- **๓.** ให้ระบุ "ข้อเสนอ" โดยเสนอเพื่อ "มีหนังสือแจ้งผลการพิจารณาให้หน่วยงานต้นเรื่องทราบ"
</หลักการร่างสำหรับเจตนา>
"""
    elif reply_intent == "ตอบรับทราบ":
        system_prompt_addendum = """
<หลักการร่างสำหรับเจตนา "ตอบรับทราบ">
- **๒.** ให้แจ้งว่าหน่วยงานของเรา ({our_department_name}) ได้รับเรื่องไว้เรียบร้อยแล้ว
- **๓.** ให้ระบุขั้นตอนที่จะดำเนินการต่อไป เช่น "จะนำเรียนผู้บังคับบัญชาเพื่อพิจารณาสั่งการต่อไป" หรือ "จะแจ้งผลให้ทราบอีกครั้ง" และลงท้ายด้วย "จึงเรียนมาเพื่อโปรดทราบ"
</หลักการร่างสำหรับเจตนา>
"""
    elif reply_intent == "ส่งต่อเรื่อง/ประสานงาน":
        system_prompt_addendum = """
<หลักการร่างสำหรับเจตนา "ส่งต่อเรื่อง/ประสานงาน">
- **๒.** ให้อธิบายว่าหน่วยงานของเรา ({our_department_name}) ได้พิจารณาแล้ว และเห็นควรส่งเรื่องต่อให้หน่วยงานที่มีอำนาจหน้าที่โดยตรง
- **๓.** ให้ระบุ "ข้อเสนอ" โดยเสนอเพื่อ "ส่งเรื่องให้ [ระบุชื่อหน่วยงานที่เกี่ยวข้อง] พิจารณาดำเนินการในส่วนที่เกี่ยวข้องต่อไป" และอาจเสนอให้มีหนังสือแจ้งหน่วยงานต้นเรื่องทราบด้วย
</หลักการร่างสำหรับเจตนา>
"""
    else: 
        return (None, "เกิดข้อผิดพลาด: ไม่พบ Template สำหรับเจตนานี้")

    # User Prompt 
    original_doc_context_parts = ["--- ข้อมูลจากหนังสือต้นเรื่อง (สำหรับใช้อ้างอิง) ---"]
    for key, value in extracted_info.items():
        if value and isinstance(value, (str, list)): 
            original_doc_context_parts.append(f"- {key}: {value}")
    original_doc_context = "\n".join(original_doc_context_parts)

    user_prompt = f"""
{original_doc_context}

---
ข้อมูลและคำสั่งสำหรับการตอบกลับ:
- ข้อ ๑ ที่ต้องใช้ขึ้นต้น: {user_provided_opening}
- เจตนาการตอบกลับ: {reply_intent}
- หน่วยงานของเรา (ผู้ตอบ): {our_department_name}

---
คำสั่ง:
จงทำหน้าที่เสมียนเอกอัจฉริยะ ร่างเนื้อหาส่วนที่เหลือ (เริ่มต้นจาก ๒.) ต่อจาก "ข้อ ๑" ที่ให้มาให้สมบูรณ์และเป็นทางการที่สุด โดยใช้ข้อมูลทั้งหมดที่ให้มาประกอบการพิจารณา และปฏิบัติตาม <หลักการร่างสำหรับเจตนา> และ <กฎการแสดงผล> อย่างเคร่งครัด
"""

    final_system_prompt = system_prompt_base + system_prompt_addendum
    return final_system_prompt, user_prompt



def replySec234_generation(client, extracted_info: dict, original_doc_type: str, reply_intent: str, relevant_internal_data: dict = None):
    if not OLLAMA_AVAILABLE:
        return "ระบบ AI (Ollama) ไม่พร้อมใช้งาน กรุณาตรวจสอบการเชื่อมต่อ"

    if not extracted_info:
        return "เกิดข้อผิดพลาด: ไม่พบข้อมูลที่สกัดจากหนังสือรับเพื่อใช้ในการสร้างการตอบกลับ"

    if not extracted_info.get("user_provided_opening_paragraph", "").strip():
        return "เกิดข้อผิดพลาด: ไม่ได้รับ 'ข้อ ๑' ที่ยืนยันแล้วจากผู้ใช้เพื่อใช้ในการร่างเนื้อหาต่อ"

    system_prompt, user_prompt = _build_prompts_for_intent(reply_intent, extracted_info, relevant_internal_data)

    if not system_prompt:
        return user_prompt

    try:
        response = client.chat(
            model= LLM_MODEL,
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            options={
                'temperature': 0.1,
                'num_predict': 2048,
                'top_p': 0.7,
                'repetition_penalty': 1.1
            }
        )
        raw_response = response['message']['content'].strip()
        
        if raw_response.startswith("```") and raw_response.endswith("```"):
            cleaned_response = '\n'.join(raw_response.split('\n')[1:-1]).strip()
        else:
            cleaned_response = raw_response

        cleaned_response = re.sub(r'<\s*[/]?[^>]+>', '', cleaned_response)
        our_department_name = (relevant_internal_data or {}).get("our_department_name", "")
        cleaned_response = cleaned_response.replace('{{our_department_name}}', our_department_name)
        cleaned_response = '\n'.join([line.strip() for line in cleaned_response.splitlines() if line.strip()])
            
        return cleaned_response
        
    except Exception as e:
        print(f"ERROR [replySec234_generation]: Exception - {e}")
        return f"เกิดข้อผิดพลาดในการเรียกใช้ AI เพื่อสร้างเนื้อหาตอบกลับ: {e}"

# Page 3 : Chat bot  

import qdrant_client
from sentence_transformers import SentenceTransformer
import numpy as np
from numpy.linalg import norm

SYSTEM_USAGE_KNOWLEDGE = """
    # คู่มือการใช้งานระบบสร้างเอกสารราชการอัจฉริยะ

    ## ภาพรวมของระบบ
    ระบบนี้มี 3 เมนูหลัก:
    1.  **หน้าแรก:** แนะนำภาพรวมและวัตถุประสงค์ของระบบ
    2.  **ร่างหนังสือราชการ:** สำหรับสร้างเนื้อหาเอกสารจากข้อความภาษาพูด
    3.  **สร้างหนังสือตอบกลับ:** สำหรับอัปโหลดหนังสือรับ (PDF) และสร้างร่างหนังสือตอบกลับ

    ## วิธีการใช้งานเมนู "ร่างหนังสือราชการ" (✍️)
    1.  **ขั้นตอนที่ 1: ใส่เนื้อหาที่ต้องการ:** ผู้ใช้พิมพ์ข้อความที่ต้องการจะสื่อสารด้วยภาษาปกติในช่องข้อความใหญ่
    2.  **ขั้นตอนที่ 2: ตั้งค่า:**
        - เลือก "ประเภทเอกสาร" ที่ต้องการ: "กระดาษข่าวร่วม (ทท.)" หรือ "บันทึกข้อความ"
        - หากเลือก "บันทึกข้อความ" จะมีตัวเลือก "คำขึ้นต้น" (เรียน/เสนอ) ให้เลือกด้วย
        - เลือก "ระดับความเป็นทางการ" ที่ต้องการ
    3.  **กดปุ่ม "✨ แปลงเป็นภาษาราชการ":** AI จะทำการประมวลผลและแสดงผลลัพธ์ในช่องข้อความด้านล่าง
    4.  **การจัดการผลลัพธ์:** ผู้ใช้สามารถคัดลอกเนื้อหา หรือกดปุ่ม "✏️ แก้ไขเนื้อหา" เพื่อปรับแก้ข้อความได้

    ## วิธีการใช้งานเมนู "สร้างหนังสือตอบกลับ" หรือ "ทำหนังสือตอบกลับ" (📬)
    เป็นกระบวนการที่มีหลายขั้นตอนที่สุด:
    1.  **ขั้นตอนที่ 1: อัปโหลดไฟล์:** ผู้ใช้อัปโหลดไฟล์ PDF ของ "หนังสือรับ" ที่ต้องการจะตอบกลับ
    2.  **ขั้นตอนที่ 1.1: ระบุประเภทเอกสาร:** เลือกว่าเอกสารที่อัปโหลดเป็น "บันทึกข้อความ" หรือ "กระดาษข่าวร่วม (ทท.)"
    3.  **ขั้นตอนที่ 1.2: สกัดข้อมูล:** กดปุ่ม "📊 สกัดข้อมูล" เพื่อให้ AI อ่านและดึงข้อมูลสำคัญจากเอกสารออกมา
    4.  **ขั้นตอนที่ 1.3: ตรวจสอบและแก้ไขข้อมูล:** ระบบจะแสดงข้อมูลที่สกัดได้ในฟอร์ม ผู้ใช้สามารถแก้ไขให้ถูกต้องแล้วกด "💾 บันทึกการแก้ไขข้อมูล"
    5.  **ขั้นตอนที่ 2.1: สร้างและยืนยัน 'ข้อ ๑':**
        - กดปุ่ม "✨ สร้างตัวเลือก 'ข้อ ๑'..." เพื่อให้ AI สร้างตัวเลือกย่อหน้าแรกของหนังสือตอบกลับ
        - ผู้ใช้สามารถเลือกตัวเลือกที่ดีที่สุดจาก Radio button หรือแก้ไขข้อความใน Text Area ให้สมบูรณ์ แล้วกด "💾 บันทึกและยืนยัน 'ข้อ ๑' นี้"
    6.  **ขั้นตอนที่ 2.2: ให้ AI ช่วยร่างเนื้อหาส่วนที่เหลือ:**
        - เลือก "หน่วยงานผู้ตอบ" และ "เจตนาหลักของการตอบกลับ" (เช่น อนุมัติ, ปฏิเสธ, ตอบรับทราบ, ส่งต่อเรื่อง)
        - กดปุ่ม "🤖 ให้ AI ช่วยร่างเนื้อหาต่อ (ข้อ ๒, ๓, ...)"
    7.  **ขั้นตอนที่ 3: ตรวจสอบและนำไปใช้งาน:**
        - ระบบจะแสดงร่างหนังสือตอบกลับฉบับสมบูรณ์ ผู้ใช้สามารถแก้ไขเป็นครั้งสุดท้าย
        - สามารถดาวน์โหลดเป็นไฟล์ .docx หรือคัดลอกเนื้อหาทั้งหมดไปใช้งานได้
    """

@st.cache_resource
def load_embedding_model():
    
    print("Loading embedding model for chatbot...")

    return SentenceTransformer('intfloat/multilingual-e5-large')

@st.cache_resource
def init_qdrant_client():
    
    print("Initializing Qdrant client...")

    return qdrant_client.QdrantClient(host="qdrant", port=6333)

qdrant_cli = init_qdrant_client()
embedding_model = load_embedding_model()

def search_in_qdrant(query: str, collection_name: str, n_results: int = 5) -> str: # << เพิ่ม n_results เป็น 5
    
    try:
        query_embedding = embedding_model.encode(query)
        search_result = qdrant_cli.search(
            collection_name=collection_name,
            query_vector=query_embedding,
            limit=n_results,
            with_payload=True 
        )
        
        formatted_context = []
        if not search_result:
            return "ไม่พบข้อมูลที่เกี่ยวข้องโดยตรง"

        for i, hit in enumerate(search_result):
            payload = hit.payload
            source_info = f"ที่มา: {payload.get('source_file', 'N/A')}, หน้า: {payload.get('page_number', 'N/A')}"
            context_block = f"""
                --- ข้อมูลอ้างอิงส่วนที่ {i+1} ---
                [เนื้อหา]: {payload.get('text', '')}
                [{source_info}]
                """
            formatted_context.append(context_block.strip())
        
        return "\n".join(formatted_context)

    except Exception as e:
        print(f"Error during Qdrant search in collection '{collection_name}': {e}")
        return "เกิดข้อผิดพลาดในการค้นหาข้อมูลจากฐานความรู้"

# # --- ฟังก์ชันการค้นหาและฟังก์ชัน Router ---
# def search_in_kb(query: str, chunks: list, embeddings: np.ndarray, n_results: int = 3) -> str:
    
#     query_embedding = embedding_model.encode(query)
#     cos_similarities = np.dot(embeddings, query_embedding) / (norm(embeddings, axis=1) * norm(query_embedding))
#     top_indices = np.argsort(cos_similarities)[::-1][:n_results]
#     top_chunks = [chunks[i] for i in top_indices]
#     return "\n---\n".join(top_chunks) if top_chunks else "ไม่พบข้อมูลที่เกี่ยวข้องโดยตรง"

def query_router(client, query: str) -> str:

    router_prompt = f"""คุณคือ AI คัดแยกคำถาม ภารกิจของคุณคือวิเคราะห์คำถามของผู้ใช้แล้วตอบกลับด้วยหนึ่งในสามคำนี้เท่านั้น: "การใช้งานระบบ", "ระเบียบสารบรรณ", หรือ "ทั่วไป"

- ถ้าคำถามเกี่ยวกับการใช้งานโปรแกรม, ขั้นตอนการทำงานของระบบ, หรือมีคำว่า "ระบบ" อยู่ในประโยคคำถาม ให้ตอบว่า: การใช้งานระบบ
- ถ้าในประโยคคำถามมีคำว่า "อ้างถึงระเบียบงานสารบรรณ" "ระเบียบ" "ข้อบังคับ" หรือชื่อเอกสารเฉพาะ ให้ตอบว่า: ระเบียบสารบรรณ
- ถ้าเป็นคำถามทักทาย, ขอบคุณ, หรือนอกเหนือจากสองเรื่องข้างบน ให้ตอบว่า: ทั่วไป

คำถามของผู้ใช้: "{query}"
ประเภทของคำถามคือ:"""
    try:
        response = client.chat(
            model=LLM_MODEL,
            messages=[{'role': 'user', 'content': router_prompt}],
            options={'temperature': 0, 'num_predict': 50}
        )
        route_result = response['message']['content'].strip()

        if "การใช้งานระบบ" in route_result:
            return "การใช้งานระบบ"
        elif "ระเบียบสารบรรณ" in route_result:
            return "ระเบียบสารบรรณ"
        else:
            return "ทั่วไป"
    except Exception as e:
        print(f"Error in query_router: {e}")

        if "ระเบียบ" in query or "สารบรรณ" in query:
            return "ระเบียบสารบรรณ"
        return "การใช้งานระบบ"

    
def call_chatbot(history: list):
    if not OLLAMA_AVAILABLE:
        return "ขออภัยครับ ระบบ AI ไม่พร้อมใช้งานในขณะนี้"

    user_query = ""
    for msg in reversed(history):
        if msg['role'] == 'user':
            user_query = msg['content']
            break
            
    if not user_query:
        return "ขออภัยครับ ผมไม่เข้าใจคำถาม กรุณาลองอีกครั้ง"

#     with st.spinner("กำลังวิเคราะห์คำถาม..."):
#         route = query_router(ollama_client, user_query)
#         st.write(f"ประเภทคำถาม: {route}")

    relevant_context = ""
#     if "การใช้งานระบบ" in route:
#         with st.spinner("กำลังค้นหาข้อมูลการใช้งานระบบ..."):
#             relevant_context = search_in_qdrant(user_query, collection_name="system_usage")
#     elif "ระเบียบสารบรรณ" in route:
#         with st.spinner("กำลังค้นหาข้อมูลในระเบียบสารบรรณ..."):
#             relevant_context = search_in_qdrant(user_query, collection_name="rtarf_knowledge_base")

#     # ถ้า router จัดเป็น "ทั่วไป" หรือค้นหาไม่เจอในหมวดอื่น ให้ลองค้นหาใน "การใช้งานระบบ" เป็นทางเลือกสุดท้าย
#     if not relevant_context or "ทั่วไป" in route:
#         print("Router classified as 'General' or no context found, attempting fallback search in 'system_usage'.")
#         with st.spinner("กำลังค้นหาข้อมูลเพิ่มเติม..."):
#             fallback_context = search_in_qdrant(user_query, collection_name="system_usage")
#             # ใช้ผลลัพธ์จาก fallback หากดีกว่าอันเดิม
#             if fallback_context and "ไม่พบข้อมูล" not in fallback_context:
#                 relevant_context = fallback_context
    with st.spinner("กำลังค้นหาข้อมูลในฐานความรู้..."):
        relevant_context = search_in_qdrant(user_query, collection_name="rtarf_knowledge_base")
    
    system_prompt = f"""คุณคือ "ที่ปรึกษาอัจฉริยะด้านงานสารบรรณ" ผู้มีทักษะการวิเคราะห์และสังเคราะห์ข้อมูลขั้นสูง ภารกิจของคุณคือการให้คำตอบที่ **ละเอียด ชัดเจน และนำไปใช้งานได้จริง** โดยอ้างอิงจากข้อมูลที่ให้มาอย่างเคร่งครัด

    **กระบวนการคิดและตอบ (Chain-of-Thought):**

    1.  **วิเคราะห์คำถาม (Analyze the Query):** อ่านคำถามของผู้ใช้ (`{user_query}`) อย่างละเอียด แล้วทำความเข้าใจเจตนาที่แท้จริง ว่าผู้ใช้ต้องการทราบอะไรกันแน่?
    2.  **สแกนและเชื่อมโยงข้อมูล (Scan & Connect Context):** อ่าน "ข้อมูลที่เกี่ยวข้อง" ทั้งหมดที่ให้มา แล้วมองหา "ทุกส่วน" ที่เกี่ยวข้องกับคำถามของผู้ใช้ แม้จะอยู่คนละส่วนกันก็ตาม จากนั้นพยายามเชื่อมโยงข้อมูลเหล่านั้นเข้าด้วยกัน
    3.  **สังเคราะห์คำตอบ (Synthesize the Answer):**
        -   **สร้างคำตอบใหม่:** ห้ามคัดลอกข้อมูลที่ให้มาแบบคำต่อคำ แต่จงใช้ภาษาของตัวเองเพื่อ "สังเคราะห์" และ "เรียบเรียง" คำตอบขึ้นมาใหม่ให้เข้าใจง่าย
        -   **ตอบให้ครบทุกประเด็น:** ตรวจสอบให้แน่ใจว่าคำตอบของคุณครอบคลุมทุกแง่มุมของคำถาม
        -   **ยกตัวอย่าง (ถ้าเป็นไปได้):** หากข้อมูลที่เกี่ยวข้องมีตัวอย่างประกอบ ให้ยกตัวอย่างนั้นมาเพื่อเพิ่มความชัดเจน
        -   **จัดรูปแบบให้อ่านง่าย:** หากคำตอบมีหลายขั้นตอนหรือหลายหัวข้อ ให้ใช้ Markdown (เช่น `-` สำหรับ bullet points หรือ `1.` สำหรับรายการ) เพื่อให้อ่านง่าย

    **กฎสำคัญ:**
    -   **ยึดตามข้อมูลเท่านั้น:** คำตอบทั้งหมดต้องมาจาก "ข้อมูลที่เกี่ยวข้อง" ที่ให้มา ห้ามเพิ่มเติมข้อมูลหรือความคิดเห็นส่วนตัวที่ไม่มีในแหล่งอ้างอิง
    -   **ยอมรับเมื่อไม่รู้:** หากวิเคราะห์แล้วพบว่า "ข้อมูลที่เกี่ยวข้อง" ไม่มีข้อมูลที่ตอบคำถามได้เลย ให้ตอบอย่างสุภาพว่า "ขออภัยครับ ผมไม่พบข้อมูลที่ชัดเจนเกี่ยวกับเรื่อง '{user_query}' ในฐานข้อมูลที่มีอยู่ครับ"
    """
    
    user_prompt_with_context = f"""
    --- ข้อมูลที่เกี่ยวข้อง ---
    {relevant_context}
    ---
    คำถามของผู้ใช้: "{user_query}"
    ---
    คำสั่ง: โปรดปฏิบัติตาม **กระบวนการคิดและตอบ (Chain-of-Thought)** ที่ระบุไว้ในบทบาทของคุณ เพื่อสร้างคำตอบที่ดีที่สุดสำหรับคำถามของผู้ใช้
    """
    
    try:
        response = ollama_client.chat(
            model=LLM_MODEL,
            messages=[
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt_with_context}
            ],
            options = {
                'temperature': 0.3, 
                'top_p': 0.9,   
                'num_predict': 4096
            }
        )
        return response['message']['content']
    except Exception as e:
        print(f"Error during chatbot call: {e}")
        return f"เกิดข้อผิดพลาดในการเชื่อมต่อกับ AI: {e}"
